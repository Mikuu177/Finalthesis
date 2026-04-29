from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import shutil, re, json, zipfile
from PIL import Image, ImageOps

ROOT = Path('/Volumes/Mikuu-ultra/Git/Finalthesis')
TEMPLATE = Path('/Volumes/Mikuu-ultra/下载/finalReportTemplateWord(1).docx')
OUT = ROOT / 'final_report_tram_prompt_routing_template.docx'
FIGDIR = ROOT / 'figures' / 'template_report'
REDESIGN_DIR = ROOT / 'figures' / 'thesis_redesign'
FIGDIR.mkdir(parents=True, exist_ok=True)

BLUE = '#2f5f8f'
DARK = '#222222'
MID = '#666666'
LIGHT = '#f4f6f8'
BORDER = '#c7d0da'

# ---------- Figure helpers ----------
def save_flow(path, title, nodes, bottom_note=None, figsize=(12, 3.2)):
    fig, ax = plt.subplots(figsize=figsize, dpi=300)
    ax.axis('off')
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.text(0.02, 0.91, title, fontsize=15, weight='bold', color=DARK, ha='left')
    n = len(nodes)
    w = 0.82 / n
    x0 = 0.04
    y = 0.45
    h = 0.22
    for i, node in enumerate(nodes):
        x = x0 + i * (0.9 / n)
        color = '#ffffff' if i not in (1, len(nodes)-2) else '#eaf1f8'
        box = FancyBboxPatch((x, y), w, h, boxstyle='round,pad=0.018,rounding_size=0.02',
                             linewidth=1.1, edgecolor=BLUE if i in (1, len(nodes)-2) else BORDER,
                             facecolor=color)
        ax.add_patch(box)
        ax.text(x + w/2, y + h/2, node, fontsize=9.2, color=DARK, ha='center', va='center', wrap=True)
        if i < n - 1:
            ax.add_patch(FancyArrowPatch((x+w+0.01, y+h/2), (x+0.9/n-0.01, y+h/2),
                                         arrowstyle='->', mutation_scale=12, linewidth=1, color=MID))
    if bottom_note:
        ax.text(0.5, 0.17, bottom_note, fontsize=9.5, color=MID, ha='center')
    fig.savefig(path, bbox_inches='tight', facecolor='white')
    plt.close(fig)

# Fig 1 closed loop
save_flow(FIGDIR/'fig01_closed_loop.png', 'Closed Research Loop', [
    'TRAM temporal QA problem',
    'Preliminary model exploration',
    'Observed category/model differences',
    'Train lightweight classifier',
    'Prompt routing + fallback',
    'Corrected evaluation',
    'Error diagnosis'
], 'The classifier is introduced as a response to observed task heterogeneity, not as an isolated module.', figsize=(13.5,3.3))

# Fig 2 preliminary exploration
fig, ax = plt.subplots(figsize=(8.6,4.5), dpi=300)
labels = ['DeepSeek\nearly run', 'Blind\nmulti-model', 'Hour24\nDeepSeek', 'Hour24\nRouter', 'TZ\nDeepSeek', 'TZ\nRouter']
vals = [0.9100, 0.8767, 0.9400, 0.9800, 0.3200, 0.4000]
colors = ['#8796a8', '#b7c0ca', '#8796a8', BLUE, '#8796a8', BLUE]
ax.bar(labels, vals, color=colors, edgecolor='#333333', linewidth=0.5)
ax.set_ylim(0,1.05)
ax.set_ylabel('Accuracy')
ax.set_title('Preliminary Exploration: Blind Multi-Model Use Was Not Enough', loc='left', weight='bold')
for i,v in enumerate(vals):
    ax.text(i, v+0.025, f'{v:.4f}', ha='center', fontsize=8.5)
ax.spines[['top','right']].set_visible(False)
ax.grid(axis='y', alpha=0.22)
fig.tight_layout()
fig.savefig(FIGDIR/'fig02_preliminary_exploration.png', bbox_inches='tight', facecolor='white')
plt.close(fig)

# Fig 3 classifier data processing
save_flow(FIGDIR/'fig03_classifier_processing.png', 'Classifier Input Processing', [
    'Raw question text\n+ optional context',
    'TF-IDF vectorizer\nunigrams + bigrams',
    'Sparse feature vector\nmax 50,000 features',
    'Logistic Regression\nmulticlass classifier',
    'Category probabilities',
    'Confidence = max probability',
    'Prompt route decision'
], 'The classifier predicts the task type; it does not generate the final answer.', figsize=(13.5,3.3))

# Reuse/redraw existing core figures and real system evidence screenshots.
copy_map = {
    'fig04_training_pipeline.png': ROOT/'figures/MODEL-Training-Pipeline.png',
    'fig05_training_cache.png': ROOT/'figures/final_en/fig04_training_cache.png',
    'fig07_main_results.png': ROOT/'figures/final_en/fig06_main_results.png',
    'fig08_category_results.png': ROOT/'figures/final_en/fig07_category_results.png',
    'fig09_remaining_errors.png': ROOT/'figures/final_en/fig08_remaining_errors.png',
    'fig10_problem_types_screenshot.png': ROOT/'figures/S2.png',
    'fig12_runtime_pipeline.png': ROOT/'figures/MODEL-System-Pipeline.png',
    'fig13_evidence_panel.png': ROOT/'figures/S14.png',
    'fig14_workflow_table_screenshot.png': ROOT/'figures/S7.png',
    'fig15_scoring_policy_screenshot.png': ROOT/'figures/S6.png',
}
for name, src in copy_map.items():
    shutil.copyfile(src, FIGDIR/name)

# Override concept/method figures with user-generated redesigned GPT images when available.
redesign_overrides = {
    'fig01_closed_loop.png': 'fig01_closed_loop_research_logic.png',
    'fig03_classifier_processing.png': 'fig02_classifier_input_processing.png',
    'fig04_training_pipeline.png': 'fig04_training_pipeline.png',
    'fig05_training_cache.png': 'fig05_training_cache_reproducibility.png',
    'fig12_runtime_pipeline.png': 'fig06_runtime_pipeline.png',
    'fig18_future_work_roadmap.png': 'fig07_future_work_roadmap.png',
}
for target, source_name in redesign_overrides.items():
    src = REDESIGN_DIR / source_name
    if src.exists():
        shutil.copyfile(src, FIGDIR / target)

# Classifier training results are reported as a formal table in Chapter 4.
# No classifier-result figure is generated, to avoid duplicating Table 4.1.

# Optional classifier architecture override. If the redesigned architecture is absent,
# do not fall back to the old figure to avoid duplicate/inconsistent diagrams.
arch_override = REDESIGN_DIR / 'fig03_classifier_architecture.png'
if arch_override.exists():
    shutil.copyfile(arch_override, FIGDIR / 'fig11_classifier_architecture.png')
else:
    old_arch = FIGDIR / 'fig11_classifier_architecture.png'
    if old_arch.exists():
        old_arch.unlink()


def compose_vertical(output, inputs, title=None, gap=22, pad=24):
    images = [Image.open(p).convert('RGB') for p in inputs]
    max_w = max(im.width for im in images)
    title_h = 70 if title else 0
    total_h = title_h + pad * 2 + sum(im.height for im in images) + gap * (len(images) - 1)
    canvas = Image.new('RGB', (max_w + pad * 2, total_h), 'white')
    y = pad
    if title:
        from PIL import ImageDraw
        d = ImageDraw.Draw(canvas)
        d.text((pad, y), title, fill=(30, 30, 30))
        y += title_h
    for im in images:
        x = pad + (max_w - im.width) // 2
        # Add a subtle border so screenshots remain visually separated in Word.
        bordered = ImageOps.expand(im, border=2, fill=(210, 218, 226))
        canvas.paste(bordered, (x, y))
        y += bordered.height + gap
    canvas.save(output)

compose_vertical(
    FIGDIR/'fig16_corrected_audit_examples.png',
    [ROOT/'figures/S8-1.png', ROOT/'figures/S8-2.png'],
    title='Corrected-match audit examples'
)
compose_vertical(
    FIGDIR/'fig17_remaining_error_samples.png',
    [ROOT/'figures/S10-1.png', ROOT/'figures/S10-2.png', ROOT/'figures/S10-3.png'],
    title='Remaining corrected error samples'
)

# ---------- Doc helpers ----------
def delete_all_content(doc):
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith('sectPr'):
            continue
        body.remove(child)

def set_black(run):
    run.font.color.rgb = RGBColor(0,0,0)

def add_p(doc, text='', style='Normal', align=None, bold=False, italic=False):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        set_black(r)
    return p

def add_heading(doc, text, level=1):
    style = 'Heading 1' if level == 1 else 'Heading 2' if level == 2 else 'Heading 3'
    p = add_p(doc, text, style=style)
    return p

def add_caption(doc, text, kind='figure'):
    style = 'figure caption' if kind == 'figure' else 'table caption'
    p = add_p(doc, text, style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if kind == 'figure' else WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_table(doc, caption, headers, rows):
    # The provided final report template places table captions below tables.
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i,h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True; set_black(r)
    for row in rows:
        cells = table.add_row().cells
        for i,v in enumerate(row):
            cells[i].text = str(v)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                for r in p.runs:
                    set_black(r)
    add_caption(doc, caption, 'table')
    add_p(doc)
    return table

def add_figure(doc, path, caption, width=6.2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(path), width=Inches(width))
    add_caption(doc, caption, 'figure')

def add_toc_placeholder(doc):
    add_heading(doc, 'Table of Contents', 1)
    p = add_p(doc, 'Right-click this area in Word/WPS and choose Update Field to generate the table of contents.', style='Normal')
    # Insert a real TOC field for Word where supported.
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'TOC \\o "1-3" \\h \\z \\u')
    p._p.append(fld)

# ---------- Build document ----------
doc = Document(str(TEMPLATE))
delete_all_content(doc)
# blank stale headers/footers to avoid template page-number artifacts being frozen
for sec in doc.sections:
    for hf in [sec.header, sec.footer]:
        for p in hf.paragraphs:
            p.text = ''

# Cover
add_p(doc, 'Final Report', 'Title1', WD_ALIGN_PARAGRAPH.CENTER)
add_p(doc, 'Classifier-Driven Prompt Routing for Complex Temporal Reasoning on TRAM', 'Title1', WD_ALIGN_PARAGRAPH.CENTER)
add_p(doc, 'Nie Wenhao', 'Subtitle1', WD_ALIGN_PARAGRAPH.CENTER)
add_p(doc, 'Supervisor: Jia Zhen', 'Subtitle1', WD_ALIGN_PARAGRAPH.CENTER)
add_p(doc, 'Submitted in accordance with the requirements for the degree of <Name of Degree>', 'centred')
add_p(doc, '<Session>', 'centred')
add_p(doc, '<Module code and name>', 'centred')
doc.add_page_break()

add_p(doc, 'The candidate confirms that the following have been submitted:', 'Normal')
add_table(doc, 'Table 0.1 Submitted materials.', ['Items','Format','Recipient(s) and Date'], [
    ['Final Report', 'DOCX/PDF file', 'To be submitted to the required university platform'],
    ['Source code and experiment artifacts', 'Project repository / local archive', 'Available for inspection'],
    ['Generated figures and analysis outputs', 'PNG/JSON/CSV files', 'Embedded or referenced in the report'],
])
add_p(doc, 'The candidate confirms that the work submitted is their own and that appropriate credit has been given where reference has been made to the work of others.', 'Footer')
add_p(doc, 'I understand that failure to attribute material obtained from another source may be considered plagiarism.', 'Footer')
add_p(doc, '(Signature of student)', 'Footer')
add_p(doc, '© <Year of Submission> The University of Leeds and Nie Wenhao', 'Normal')
doc.add_page_break()

# Summary
add_heading(doc, 'Summary', 1)
summary_paras = [
    'Temporal reasoning is a central capability for practical language understanding, but it is not a single uniform skill. Tasks such as date computation, hour adjustment, time-zone conversion, and year shifting require different operations even though they all appear to involve time. This project studies complex temporal question answering on a TRAM-derived setting and investigates whether a lightweight classifier can make prompt selection more task-aware.',
    'The project began with preliminary model and workflow exploration. A strong single-model run reached 0.9100 accuracy on an early 100-sample setting, while a blind multi-model workflow reached 0.8767 despite using three calls per query. Category-specific experiments also showed different behavior: on Hour Adjustment (24h), profile-based routing improved from 0.9400 to 0.9800, while on Time Zone Conversion it improved from 0.3200 to 0.4000. These observations motivated a more controlled approach: instead of blindly adding models, the system should first recognize the temporal problem type and then choose a suitable prompt.',
    'The final method trains a small TF-IDF plus Logistic Regression classifier to predict five temporal categories: Date Computation, Hour Adjustment (24h), Month Shift, Time Zone Conversion, and Year Shift. The classifier is trained on 6709 samples, evaluated on a 958-sample development split, and tested on an independent 1918-sample split. It reaches 1.0000 accuracy and 1.0000 macro-F1 on both development and independent test splits. The classifier output is used to route each question to a category-specific prompt, with a fallback prompt when confidence is lower than 0.95.',
    'Under corrected evaluation ruleset v1.1, which normalizes equivalent answer formats only at scoring time, the four final workflows achieve the following corrected accuracies: Fixed Prompt 0.7575, CoT Prompt 0.7700, Classifier Router 0.7575, and Classifier Router + Fallback 0.7750. The best observed workflow is therefore Classifier Router + Fallback, without increasing the number of LLM calls. Error analysis shows that 376 corrected errors remain, mainly in Date Computation, Time Zone Conversion, and Hour Adjustment (24h). The project therefore contributes not only a routing workflow, but also a clearer measurement and diagnosis framework for temporal reasoning.',
]
for x in summary_paras: add_p(doc, x)

doc.add_page_break()
add_heading(doc, 'Acknowledgements', 1)
add_p(doc, 'I would like to thank my supervisor, Jia Zhen, for guidance and feedback during the project. I also acknowledge the authors of TRAM and the open-source software communities whose tools made the implementation and evaluation possible.')

doc.add_page_break()
add_toc_placeholder(doc)
doc.add_section(WD_SECTION.NEW_PAGE)

# Chapter 1
add_heading(doc, 'Chapter 1\nIntroduction and Background Research', 1)
add_heading(doc, '1.1 Introduction', 2)
for x in [
    'Large language models have shown strong general language ability, but temporal reasoning remains challenging because it combines natural language interpretation with symbolic operations. Temporal reasoning benchmarks show that time-related tasks differ across arithmetic, event, commonsense, and question-answering settings (Wang and Zhao, 2024; Chu et al., 2024). A model may need to add days across a month boundary, convert time zones, adjust minutes with carry or borrow, or identify a shifted year. These operations differ structurally, so a single prompt is unlikely to be equally effective for every temporal subtask.',
    'This report focuses on the TRAM temporal reasoning benchmark and develops a classifier-driven prompt routing system. TRAM provides a broad benchmark for temporal reasoning in large language models, while TimeBench further emphasizes category-level differences in temporal reasoning ability (Wang and Zhao, 2024; Chu et al., 2024). The central question is whether a small, trainable classifier can identify the temporal category of a question and use that category to select a more suitable prompt for a large language model. The classifier is not used to produce the final answer; it acts as a task-recognition and routing layer before the LLM.',
]: add_p(doc,x)

add_figure(doc, FIGDIR/'fig01_closed_loop.png', 'Figure 1.1 Closed research loop from preliminary model exploration to classifier-driven routing and error diagnosis.', 6.4)

add_heading(doc, '1.2 Preliminary Multi-Model Exploration', 2)
for x in [
    'The classifier was introduced after an exploratory stage rather than as an isolated design choice. Early experiments compared a strong single-model baseline, a blind multi-model workflow, and profile-based routing on selected temporal categories. This motivation is consistent with prior work on routing, which treats selection as a way to improve performance and cost trade-offs rather than simply using more model calls (Hu et al., 2024). The results showed two important patterns. First, blindly calling more models was not automatically better. Second, model and workflow performance varied across temporal categories.',
    'In the early 100-sample setting, the DeepSeek single-model baseline reached 0.9100 accuracy, while a blind multi-model workflow reached 0.8767 and required three calls per query. Later category-specific experiments showed why task awareness mattered: on Hour Adjustment (24h), router-only improved over the strong single model, while on Time Zone Conversion the same general idea also improved accuracy but remained much harder overall.',
]: add_p(doc,x)
add_figure(doc, FIGDIR/'fig02_preliminary_exploration.png', 'Figure 1.2 Preliminary exploration showing that blind multi-model calling was not sufficient and that category-sensitive routing was promising.', 5.9)
add_table(doc, 'Table 1.1 Preliminary exploration summary.', ['Setting','Compared systems','Key result','Interpretation'], [
    ['Early 100-sample run', 'DeepSeek vs blind multi-model', '0.9100 vs 0.8767', 'More model calls did not guarantee better accuracy.'],
    ['Hour24 split', 'DeepSeek vs router-only', '0.9400 vs 0.9800', 'Category-aware routing helped on a structured temporal task.'],
    ['Time Zone split', 'DeepSeek vs router-only', '0.3200 vs 0.4000', 'Routing helped, but the category remained difficult.'],
])
add_figure(doc, FIGDIR/'fig10_problem_types_screenshot.png', 'Figure 1.3 TRAM temporal problem type examples displayed in the project workbench.', 6.2)
for x in [
    'Temporal question answering has also been studied before the recent LLM benchmark wave. Jia et al. (2018a) introduced TempQuestions as a benchmark for temporal question answering, while Jia et al. (2018b) proposed TEQUILA, which detects temporal intent and reasons over temporal constraints on top of KB-QA engines. More recently, Jia, Christmann and Weikum (2024) emphasized faithful temporal question answering over heterogeneous sources, where temporal constraints need to be handled explicitly and supported by evidence. These works motivate the view that temporal QA benefits from explicit task recognition and constraint-aware processing.',
]: add_p(doc,x)

add_heading(doc, '1.3 Motivation for a Lightweight Classifier', 2)
for x in [
    'The preliminary experiments suggested that category information was useful. Prior temporal QA research also shows that temporal questions require explicit recognition of temporal intent, temporal constraints, and evidence conditions (Jia et al., 2018a; Jia et al., 2018b; Jia, Christmann and Weikum, 2024). However, a system cannot rely on manually supplied category labels when answering a new user question. This creates the need for an automatic task recognizer. The project therefore introduces a lightweight classifier as the system’s own trainable component.',
    'This design keeps the LLM unchanged and places a small, auditable decision layer before it. The classifier predicts the temporal task type and a confidence score. The prompt router then uses this prediction to choose either a category-specific prompt or, if confidence is low, a conservative fallback prompt. This is related to selective classification, where uncertain predictions can be rejected or handled conservatively to reduce downstream errors (Geifman and El-Yaniv, 2017).',
]: add_p(doc,x)

add_heading(doc, '1.4 Research Objectives', 2)
for x in [
    'The first objective is to train and evaluate a lightweight classifier for temporal task recognition. The second objective is to integrate the classifier into a prompt routing workflow. The third objective is to compare the routed workflow against fixed-prompt and chain-of-thought baselines. The fourth objective is to separate true reasoning errors from answer-format artifacts through corrected evaluation.',
]: add_p(doc,x)

# Chapter 2
add_heading(doc, 'Chapter 2\nMethodology', 1)
add_heading(doc, '2.1 Dataset and Strict Split', 2)
for x in [
    'The project uses a TRAM-derived temporal reasoning dataset and focuses on five categories: Date Computation, Hour Adjustment (24h), Month Shift, Time Zone Conversion, and Year Shift. TRAM is designed to cover multiple temporal aspects, including arithmetic and duration-related reasoning, rather than treating temporal reasoning as one uniform task (Wang and Zhao, 2024). These categories were selected because they represent different forms of temporal arithmetic and provide a clear setting for task-aware routing.',
    'The classifier is trained and evaluated with a strict split. The training set contains 6709 samples, the development set contains 958 samples, and the independent test set contains 1918 samples. The final LLM workflow comparison uses a frozen 400-sample evaluation slice. This separation is important because classifier performance and final QA performance are related but distinct measurements.',
]: add_p(doc,x)

add_heading(doc, '2.2 Classifier Input Processing and Architecture', 2)
for x in [
    'The classifier receives the raw question text and optional context. These fields are concatenated into a single text input. The TF-IDF vectorizer then converts the text into unigram and bigram features using `ngram_range=(1, 2)`, `min_df=2`, and `max_features=50000`. TF-IDF is a standard sparse text representation for information retrieval and text classification (Manning, Raghavan and Schütze, 2008). A multiclass Logistic Regression model maps the sparse feature vector to category probabilities, implemented through the scikit-learn machine-learning library (Pedregosa et al., 2011).',
    'The predicted category is the class with the highest probability. Confidence is defined as the maximum predicted probability. This confidence value is later used by the fallback mechanism. The classifier therefore produces not only a label, but also a risk signal for routing. Because model probabilities are not always well calibrated, confidence is treated here as an engineering risk proxy rather than as a guaranteed estimate of correctness (Guo et al., 2017).',
]: add_p(doc,x)
add_figure(doc, FIGDIR/'fig03_classifier_processing.png', 'Figure 2.1 Classifier input processing from raw question text to prompt route decision.', 6.4)
if (FIGDIR/'fig11_classifier_architecture.png').exists():
    add_figure(doc, FIGDIR/'fig11_classifier_architecture.png', 'Figure 2.2 Lightweight classifier architecture showing TF-IDF features, Logistic Regression, category probabilities, and routing output.', 6.2)
    training_fig_no = '2.3'
    cache_fig_no = '2.4'
else:
    training_fig_no = '2.2'
    cache_fig_no = '2.3'

add_heading(doc, '2.3 Classifier Training Pipeline', 2)
for x in [
    'Training follows a supervised text-classification pipeline. The vectorizer is fitted only on the training data. Logistic Regression is trained on the resulting TF-IDF vectors. The trained pipeline is saved as a joblib artifact and evaluated on development and independent test splits.',
    'A cache mechanism is used to make training reproducible. The cache key is based on the split and training configuration. If the same configuration has already generated a model and report, the system can reuse the artifact; otherwise, it retrains and records a new report.',
]: add_p(doc,x)
add_figure(doc, FIGDIR/'fig04_training_pipeline.png', f'Figure {training_fig_no} Lightweight classifier training pipeline.', 6.2)
add_figure(doc, FIGDIR/'fig05_training_cache.png', f'Figure {cache_fig_no} Training cache and reproducibility mechanism.', 5.8)

add_heading(doc, '2.4 Prompt Routing and Fallback', 2)
for x in [
    'The prompt router uses the classifier output to choose the inference prompt. If confidence is at least 0.95, the system uses the category-specific prompt for the predicted temporal type. If confidence is below 0.95, the system uses a conservative fallback prompt. The fallback decision happens before the LLM call, so it does not increase calls per query. The design is influenced by selective prediction, but applies the idea to prompt choice rather than to final answer abstention (Geifman and El-Yaniv, 2017).',
    'This mechanism is designed as risk control. A wrong category prediction can send a question to an unsuitable prompt, so low-confidence cases should not be routed aggressively. The fallback policy provides a safer path while preserving the single-call constraint. Unlike chain-of-thought prompting, which changes the reasoning style of the prompt (Wei et al., 2022; Kojima et al., 2022), this project changes which prompt is selected for each temporal category.',
]: add_p(doc,x)

add_heading(doc, '2.5 Evaluation Normalization', 2)
for x in [
    'The project reports two evaluation views. Base evaluation uses strict string matching. Corrected evaluation uses ruleset v1.1, a category-aware normalization layer. This layer normalizes equivalent answer formats at scoring time without changing the original gold labels or model outputs.',
    'The corrected metric is necessary because many apparent errors were answer-format mismatches. For example, equivalent date formats, `H:MM` and `HH:MM`, and compact time-zone strings can represent the same answer but fail strict string matching. Corrected evaluation should therefore be interpreted as measurement correction, not as a new model improvement.',
]: add_p(doc,x)
add_table(doc, 'Table 2.1 Category-aware normalization principles in ruleset v1.1.', ['Category','Normalization principle'], [
    ['Date Computation', 'Normalize equivalent date-like formats such as MM-DD-YYYY, M/D/YYYY, and YYYY-MM-DD.'],
    ['Hour Adjustment (24h)', 'Treat H:MM and HH:MM as equivalent.'],
    ['Time Zone Conversion', 'Convert compact AM/PM time-zone strings to HH:MM when possible.'],
    ['Month Shift', 'Normalize month names, abbreviations, and month-level formats.'],
    ['Year Shift', 'Normalize integer year strings.'],
])

# Chapter 3
add_heading(doc, 'Chapter 3\nSystem Implementation', 1)
add_heading(doc, '3.1 Runtime Workflow', 2)
for x in [
    'At runtime, a user question first enters the classifier. The classifier returns a predicted category, a probability distribution, and a confidence score. The router then selects the prompt. The selected prompt and question are sent to the serving LLM, which produces the final answer. The answer is parsed and evaluated using the corrected scoring rules.',
    'This pipeline creates a clear separation of responsibilities. The classifier performs task recognition. The prompt router performs decision control. The LLM performs answer generation. The evaluator determines whether the final answer is equivalent to the gold answer.',
]: add_p(doc,x)
add_figure(doc, FIGDIR/'fig12_runtime_pipeline.png', 'Figure 3.1 Runtime pipeline for classifier-driven prompt routing with fallback and single-call LLM inference.', 6.4)

add_heading(doc, '3.2 Evidence Workbench', 2)
for x in [
    'The project also includes a local evidence workbench for demonstration and debugging. The workbench is not the research contribution itself; it is a way to inspect classifier predictions, prompt routing decisions, model outputs, and remaining error samples. This makes the system more auditable during supervision meetings and defense preparation.',
]: add_p(doc,x)
add_figure(doc, FIGDIR/'fig13_evidence_panel.png', 'Figure 3.2 Evidence panel for scoring policy, corrected summary, category-wise changes, and remaining error diagnosis.', 6.2)

add_heading(doc, '3.3 Reproducibility Artifacts', 2)
add_table(doc, 'Table 3.1 Main reproducibility artifacts.', ['Artifact','Purpose'], [
    ['classifier_router/train.jsonl', 'Classifier training split.'],
    ['classifier_router/dev.jsonl', 'Classifier development split.'],
    ['classifier_router/test.jsonl', 'Independent classifier test split.'],
    ['classifier_report.json', 'Classifier training metrics and configuration.'],
    ['eval_predictions.csv', 'Sample-level classifier predictions.'],
    ['analysis_payload_strict.json', 'Corrected workflow and category-level analysis.'],
])

# Chapter 4
add_heading(doc, 'Chapter 4\nExperiments and Results', 1)
add_heading(doc, '4.1 Evaluation Metrics', 2)
for x in [
    'The final QA workflows are evaluated with accuracy, parse rate, format compliance, latency per query, and calls per query. Classifier performance is evaluated separately with accuracy and macro-F1. Macro-F1 is included because the category distribution is imbalanced. Accuracy is also the main reporting metric in TRAM-style temporal reasoning evaluation (Wang and Zhao, 2024).',
]: add_p(doc,x)

add_heading(doc, '4.2 Classifier Training Results', 2)
for x in [
    'The classifier reaches perfect performance on the selected five-category split. On the development set, accuracy and macro-F1 are both 1.0000. On the independent test set, all 1918 samples are correctly classified. This result shows that the selected categories are highly separable at the text-classification level.',
    'This result must be interpreted carefully. It does not mean that temporal question answering is solved. The classifier only predicts task type. The final answer still depends on prompt quality and the LLM’s ability to perform the temporal computation.',
]: add_p(doc,x)
add_table(doc, 'Table 4.1 Lightweight classifier training configuration and evaluation results.', ['Stage','Item','Value'], [
    ['Data Split', 'Training samples', '6709'],
    ['Data Split', 'Development samples', '958'],
    ['Data Split', 'Independent test samples', '1918'],
    ['Feature Extraction', 'Feature extractor', 'TF-IDF'],
    ['Feature Extraction', 'N-gram range', '1-2'],
    ['Feature Extraction', 'Minimum document frequency', '2'],
    ['Feature Extraction', 'Maximum features', '50000'],
    ['Classifier', 'Model', 'Logistic Regression'],
    ['Classifier', 'Regularization', 'L2, C=1.0'],
    ['Classifier', 'Maximum iterations', '1200'],
    ['Evaluation', 'Development accuracy', '1.0000'],
    ['Evaluation', 'Development macro-F1', '1.0000'],
    ['Evaluation', 'Test accuracy', '1.0000'],
    ['Evaluation', 'Test macro-F1', '1.0000'],
    ['Evaluation', 'Test correct samples', '1918 / 1918'],
])

add_heading(doc, '4.3 Final Workflow Results', 2)
for x in [
    'The final comparison includes four workflows: Fixed Prompt, CoT Prompt, Classifier Router, and Classifier Router + Fallback. CoT Prompt is included because chain-of-thought and zero-shot reasoning prompts are strong and widely used baselines for LLM reasoning tasks (Wei et al., 2022; Kojima et al., 2022). All workflows use one LLM call per query. Under corrected evaluation, Classifier Router + Fallback is the strongest observed workflow, reaching 0.7750 accuracy.',
]: add_p(doc,x)
add_figure(doc, FIGDIR/'fig07_main_results.png', 'Figure 4.1 Corrected end-to-end workflow accuracy.', 5.8)
add_table(doc, 'Table 4.2 Workflow comparison under base and corrected evaluation.', ['Workflow','Corrected accuracy','Base accuracy','Delta','Latency (s/query)','Calls/query'], [
    ['Fixed Prompt','0.7575','0.6300','+0.1275','5.3711','1.0'],
    ['CoT Prompt','0.7700','0.6200','+0.1500','5.2518','1.0'],
    ['Classifier Router','0.7575','0.6075','+0.1500','5.2627','1.0'],
    ['Router + Fallback','0.7750','0.6250','+0.1500','5.3329','1.0'],
])
add_figure(doc, FIGDIR/'fig14_workflow_table_screenshot.png', 'Figure 4.2 Real workflow table screenshot from the analysis workbench, showing base and corrected workflow results.', 5.8)
for x in [
    'The improvement from Classifier Router to Classifier Router + Fallback is important because it shows that confidence-based risk control repairs some routing losses. The gain over CoT Prompt is small, so it should not be overstated, but it is achieved without increasing the number of LLM calls.',
]: add_p(doc,x)

add_heading(doc, '4.4 Category-wise Results', 2)
for x in [
    'Category-level results show why the overall number alone is insufficient. Date Computation improves substantially after normalization but remains difficult. Time Zone Conversion is no longer zero after corrected scoring, but it remains weak. Year Shift is already at ceiling, while Hour24 errors appear to be genuine carry and borrow reasoning failures rather than formatting artifacts.',
]: add_p(doc,x)
add_figure(doc, FIGDIR/'fig08_category_results.png', 'Figure 4.3 Category-wise corrected accuracy.', 5.4)
add_table(doc, 'Table 4.3 Category-wise corrected accuracy.', ['Category','Fixed','CoT','Router','Router + Fallback'], [
    ['Date Computation','0.7008','0.7159','0.7045','0.7235'],
    ['Hour Adjustment (24h)','0.9545','0.9773','0.8864','0.9318'],
    ['Time Zone Conversion','0.2000','0.2000','0.3000','0.3000'],
    ['Year Shift','1.0000','1.0000','1.0000','1.0000'],
])

# Chapter 5
add_heading(doc, 'Chapter 5\nError Analysis and Discussion', 1)
add_heading(doc, '5.1 Why the Original Accuracy Was Around 0.6', 2)
for x in [
    'The original base scores were around 0.6 because they used strict string matching. This was too harsh for open temporal answers. Many outputs were semantically equivalent to the gold answer but appeared in a different valid format. Corrected evaluation raises the measured performance by recognizing these equivalences.',
    'This correction does not change the model output or the original gold label. It only changes the scoring comparison. Therefore, corrected gains should be described as measurement correction, not as model improvement.',
]: add_p(doc,x)
add_figure(doc, FIGDIR/'fig15_scoring_policy_screenshot.png', 'Figure 5.1 Scoring policy screenshot showing base strict matching and corrected category-aware evaluation.', 6.0)
add_figure(doc, FIGDIR/'fig16_corrected_audit_examples.png', 'Figure 5.2 Corrected-match audit examples showing cases where strict base scoring was too harsh.', 6.2)

add_heading(doc, '5.2 Remaining Corrected Errors', 2)
for x in [
    'After ruleset v1.1, 376 errors remain. These errors are more meaningful because many format artifacts have already been removed. The distribution shows that Date Computation is the dominant bottleneck, followed by Time Zone Conversion and Hour Adjustment (24h).',
]: add_p(doc,x)
add_figure(doc, FIGDIR/'fig09_remaining_errors.png', 'Figure 5.3 Remaining corrected error distribution.', 5.6)
add_table(doc, 'Table 5.1 Remaining corrected errors.', ['Category','Remaining errors','Representative causes'], [
    ['Date Computation','305','month/day confusion, cross-month or cross-year update errors, multi-field date mistakes'],
    ['Time Zone Conversion','60','direction errors, one-hour offset errors, minute mismatch'],
    ['Hour Adjustment (24h)','11','hour and minute carry/borrow errors'],
    ['Total','376','remaining reasoning errors after normalization'],
])
add_figure(doc, FIGDIR/'fig17_remaining_error_samples.png', 'Figure 5.4 Remaining error samples after ruleset v1.1 correction, grouped by Date Computation, Time Zone Conversion, and Hour Adjustment (24h).', 6.2)

add_heading(doc, '5.3 Discussion', 2)
for x in [
    'The full loop of evidence is now clearer. Preliminary experiments showed that temporal categories behave differently and that blind multi-model use is not enough. This connects to earlier temporal QA work, where temporal intent and temporal constraints are treated as specialized structures rather than ordinary question text (Jia et al., 2018a; Jia et al., 2018b). The classifier then automated task recognition, making prompt routing possible for new inputs. The final Router + Fallback workflow became the strongest observed workflow under corrected evaluation. However, the remaining errors show that routing alone cannot solve difficult temporal computation.',
    'The main lesson is that task-aware routing is useful, but it depends on the quality of the prompt bank and the LLM’s underlying temporal reasoning ability. The classifier can select a route, but it cannot guarantee that the LLM will execute the required arithmetic correctly.',
]: add_p(doc,x)

# Chapter 6
add_heading(doc, 'Chapter 6\nConclusion and Future Work', 1)
for x in [
    'This project presents a classifier-driven prompt routing framework for complex temporal reasoning on TRAM-derived data. The work began with preliminary model exploration, which showed that temporal reasoning performance varies by category and that blind multi-model workflows are not automatically better. This motivated the introduction of a lightweight classifier as an explicit task-recognition module.',
    'The classifier uses TF-IDF features and Logistic Regression to predict temporal categories. It achieves 1.0000 accuracy and macro-F1 on both development and independent test splits, showing strong category separability. The classifier output is then used for prompt routing and confidence-based fallback. Under corrected evaluation, Classifier Router + Fallback reaches 0.7750 accuracy, the strongest observed workflow in the current experiment, while keeping one LLM call per query.',
    'Future work should follow the remaining corrected error surface. The first priority is to improve Date Computation prompts, especially for cross-month and cross-year operations. The second priority is to add stronger directional constraints for Time Zone Conversion. The third priority is to improve Hour24 carry and borrow handling. These directions are concrete because the corrected evaluation removes many format artifacts and exposes the remaining reasoning failures more clearly.',
]: add_p(doc,x)
if (FIGDIR/'fig18_future_work_roadmap.png').exists():
    add_figure(doc, FIGDIR/'fig18_future_work_roadmap.png', 'Figure 6.1 Future work roadmap following the corrected error surface.', 6.2)

# References
add_heading(doc, 'List of References', 1)
refs = [
    'Chu, Z., Chen, J., Chen, Q., Yu, W., Wang, H., Liu, M. and Qin, B. (2024) \'TimeBench: A Comprehensive Evaluation of Temporal Reasoning Abilities in Large Language Models\', Proceedings of the 62nd Annual Meeting of the Association for Computational Linguistics (Volume 1: Long Papers), pp. 1204-1228. doi: 10.18653/v1/2024.acl-long.66.',
    'Geifman, Y. and El-Yaniv, R. (2017) \'Selective Classification for Deep Neural Networks\', Advances in Neural Information Processing Systems, 30.',
    'Guo, C., Pleiss, G., Sun, Y. and Weinberger, K.Q. (2017) \'On Calibration of Modern Neural Networks\', Proceedings of the 34th International Conference on Machine Learning, PMLR 70, pp. 1321-1330.',
    'Hu, Q.J., Bieker, J., Li, X., Jiang, N., Keigwin, B., Ranganath, G., Keutzer, K. and Upadhyay, S.K. (2024) \'RouterBench: A Benchmark for Multi-LLM Routing System\', arXiv preprint arXiv:2403.12031.',
    'Jia, Z., Abujabal, A., Saha Roy, R., Strötgen, J. and Weikum, G. (2018a) \'TempQuestions: A Benchmark for Temporal Question Answering\', Companion Proceedings of The Web Conference 2018, pp. 1057-1062. doi: 10.1145/3184558.3191536.',
    'Jia, Z., Abujabal, A., Saha Roy, R., Strötgen, J. and Weikum, G. (2018b) \'TEQUILA: Temporal Question Answering over Knowledge Bases\', Proceedings of the 27th ACM International Conference on Information and Knowledge Management, pp. 1807-1810. doi: 10.1145/3269206.3269247.',
    'Jia, Z., Christmann, P. and Weikum, G. (2024) \'Faithful Temporal Question Answering over Heterogeneous Sources\', Proceedings of the ACM Web Conference 2024, pp. 2052-2063. doi: 10.1145/3589334.3645547.',
    'Kojima, T., Gu, S.S., Reid, M., Matsuo, Y. and Iwasawa, Y. (2022) \'Large Language Models are Zero-Shot Reasoners\', Advances in Neural Information Processing Systems, 35, pp. 22199-22213.',
    'Manning, C.D., Raghavan, P. and Schütze, H. (2008) Introduction to Information Retrieval. Cambridge: Cambridge University Press.',
    'Pedregosa, F. et al. (2011) \'scikit-learn: Machine Learning in Python\', Journal of Machine Learning Research, 12, pp. 2825-2830.',
    'Wang, Y. and Zhao, Y. (2024) \'TRAM: Benchmarking Temporal Reasoning for Large Language Models\', Findings of the Association for Computational Linguistics: ACL 2024, pp. 6389-6415. doi: 10.18653/v1/2024.findings-acl.382.',
    'Wei, J., Wang, X., Schuurmans, D., Bosma, M., Ichter, B., Xia, F., Chi, E., Le, Q.V. and Zhou, D. (2022) \'Chain-of-Thought Prompting Elicits Reasoning in Large Language Models\', Advances in Neural Information Processing Systems, 35, pp. 24824-24837.',
]
for r in refs: add_p(doc, r, 'hanging indent')

# Appendices
add_heading(doc, 'Appendix A Self-appraisal', 1)
add_p(doc, 'The main technical contribution of the project is the integration of a lightweight supervised classifier into a prompt-routing workflow for temporal reasoning. The work includes data splitting, classifier training, runtime routing, fallback design, corrected evaluation, and sample-level error analysis. The main limitation is that the final workflow still depends on prompt quality and LLM reasoning capability.')

add_heading(doc, 'Appendix B External Materials and Artifacts', 1)
add_table(doc, 'Table B.1 Project artifacts.', ['Artifact','Purpose'], [
    ['temporal-reasoning-benchmark/data/splits/classifier_router/train.jsonl', 'Classifier training split.'],
    ['temporal-reasoning-benchmark/data/splits/classifier_router/dev.jsonl', 'Classifier development split.'],
    ['temporal-reasoning-benchmark/data/splits/classifier_router/test.jsonl', 'Independent classifier test split.'],
    ['temporal-reasoning-benchmark/outputs/classifier_strict/classifier_report.json', 'Classifier training report.'],
    ['temporal-reasoning-benchmark/outputs/classifier_strict_eval/eval_predictions.csv', 'Independent classifier predictions.'],
    ['backend/runtime/analysis_payload_strict.json', 'Corrected workflow summary and error analysis.'],
])

# Enforce black text and font size sanity
for p in doc.paragraphs:
    for r in p.runs:
        set_black(r)
        if r.font.size is None and p.style.name == 'Normal':
            r.font.size = Pt(11)
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    set_black(r)

# Save
doc.save(str(OUT))

# validation
check = Document(str(OUT))
text = '\n'.join(p.text for p in check.paragraphs)
for t in check.tables:
    for row in t.rows:
        for cell in row.cells:
            text += '\n' + cell.text
print('OUT', OUT)
print('paragraphs', len(check.paragraphs), 'tables', len(check.tables), 'images', len(check.inline_shapes))
print('has_chinese', bool(re.search(r'[\u4e00-\u9fff]', text)))
print('bad_placeholders', any(x in text for x in ['[FIGURE]','[TABLE]','待补录']))
